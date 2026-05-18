from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("python-docx is not installed. Please install with: pip install python-docx") from exc


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main():
    out_path = Path("SharePoint_List_Write_Documentation.docx")

    doc = Document()
    doc.add_heading("SharePoint List Write Documentation", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This document summarizes which SharePoint lists the Lease File Audit app writes to, "
        "what each list stores, and whether cleanup removes those records."
    )

    doc.add_heading("Primary Lists", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "List"
    hdr[1].text = "Purpose"
    hdr[2].text = "Rows"
    hdr[3].text = "Cleared by cleanup"

    rows = [
        ("AuditRuns and AuditRuns2", "Detailed findings and bucket results", "Many (per finding/bucket)", "No"),
        ("RunDisplaySnapshots", "Portfolio, property, and lease snapshot views", "Few (per scope)", "Yes"),
        ("Audit Run Metrics", "High-level run summary statistics", "1 per run", "Yes"),
        ("ExceptionMonths", "Month-level exception tracking and resolution", "On status updates", "Yes"),
        ("LeaseTermSet, LeaseTerms, LeaseTermEvidence", "Lease term extraction and supporting evidence", "Variable", "Yes"),
        ("Innovation Use Log", "User/session activity logging", "Per event", "No"),
    ]

    for r in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

    doc.add_heading("AuditRuns and AuditRuns2 fields", level=1)
    add_bullets(
        doc,
        [
            "RunId, ResultType, PropertyId, LeaseIntervalId, ArCodeId, AuditMonth",
            "Status, Severity, FindingTitle, Variance, ExpectedTotal, ActualTotal, ImpactAmount",
            "MatchRule, FindingId, Category, Description, ExpectedValue, ActualValue, CreatedAt",
            "Optional columns when available: PropertyName, ResidentName",
        ],
    )

    doc.add_heading("RunDisplaySnapshots fields", level=1)
    add_bullets(
        doc,
        [
            "Title, SnapshotKey, RunId, ScopeType, PropertyId, LeaseIntervalId",
            "ExceptionCountStatic, UnderchargeStatic, OverchargeStatic, MatchRateStatic",
            "TotalBucketsStatic, MatchedBucketsStatic, CreatedAt",
            "Optional: PropertyNameStatic, TotalVarianceStatic, AuditedThrough",
        ],
    )

    doc.add_heading("Audit Run Metrics fields", level=1)
    add_bullets(
        doc,
        [
            "Title, RunDateTime, UploadedBy",
            "TotalScheduled, TotalActual, Matched",
            "ScheduledNotBilled, BilledNotScheduled, AmountMismatch",
            "TotalVariances, HighSeverity, MediumSeverity",
            "Properties (JSON per-property breakdown)",
        ],
    )

    doc.add_heading("ExceptionMonths fields", level=1)
    add_bullets(
        doc,
        [
            "CompositeKey, RunId, PropertyId, LeaseIntervalId, ArCodeId, AuditMonth",
            "ExceptionType, Status, FixLabel, Variance",
            "ResolvedAt, ResolvedBy, Notes",
        ],
    )

    doc.add_heading("Operational notes", level=1)
    add_bullets(
        doc,
        [
            "AuditRuns and RunDisplaySnapshots support batch writes with fallback behavior.",
            "Writes can run sync or async based on environment toggles.",
            "The cleanup script preserves AuditRuns/AuditRuns2 but clears snapshot, metrics, exception, and lease-term lists.",
        ],
    )

    doc.add_heading("Source references", level=1)
    add_bullets(
        doc,
        [
            "SHAREPOINT_WRITE_TARGETS.md",
            "SHAREPOINT_RUN_OUTPUT_MAPPING.md",
        ],
    )

    doc.save(out_path)
    print(f"Created {out_path.resolve()}")


if __name__ == "__main__":
    main()
