from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("python-docx is not installed. Install with: pip install python-docx") from exc


def _set_cell_text(cell, text):
    cell.text = text


def main():
    out_path = Path("SharePoint_Action_Read_Write_Matrix.docx")

    rows = [
        {
            "action": "Global session lifecycle logging",
            "routes": "All authenticated requests (app.before_request)",
            "reads": "None",
            "writes": "Innovation Use Log",
            "notes": "Logs Start Session / timeout End Session events.",
        },
        {
            "action": "Home page",
            "routes": "GET /",
            "reads": "LeaseFileAudit Runs library (recent runs metadata)",
            "writes": "None",
            "notes": "Session activity may also be logged by global middleware.",
        },
        {
            "action": "End session",
            "routes": "GET /end-session",
            "reads": "None",
            "writes": "Innovation Use Log",
            "notes": "Explicit End Session activity log write.",
        },
        {
            "action": "Settings",
            "routes": "GET/POST /settings",
            "reads": "None (local JSON config files)",
            "writes": "None (local JSON config files)",
            "notes": "No SharePoint list read/write in this route.",
        },
        {
            "action": "Get runs API",
            "routes": "GET /api/runs",
            "reads": "LeaseFileAudit Runs library (run metadata/list)",
            "writes": "None",
            "notes": "Used by run pickers.",
        },
        {
            "action": "Clear cache API",
            "routes": "POST /api/admin/clear-cache",
            "reads": "None",
            "writes": "None",
            "notes": "In-memory cache clear only.",
        },
        {
            "action": "Property picklist API",
            "routes": "GET /api/property-picklist",
            "reads": "None (Entrata API)",
            "writes": "None",
            "notes": "No SharePoint list usage.",
        },
        {
            "action": "Excel upload audit",
            "routes": "POST /upload",
            "reads": "May read baseline run data from LeaseFileAudit Runs library",
            "writes": "AuditRuns2; RunDisplaySnapshots; Audit Run Metrics; LeaseFileAudit Runs library; Innovation Use Log",
            "notes": "Primary audit pipeline save_run path.",
        },
        {
            "action": "Single property API audit",
            "routes": "POST /upload-api-property",
            "reads": "Entrata API; optional baseline run read from LeaseFileAudit Runs library",
            "writes": "AuditRuns2; RunDisplaySnapshots; Audit Run Metrics; LeaseFileAudit Runs library; Innovation Use Log",
            "notes": "Runs save_run and redirects to property view.",
        },
        {
            "action": "Single lease API audit",
            "routes": "POST /upload-api-lease",
            "reads": "Entrata API",
            "writes": "AuditRuns2; RunDisplaySnapshots; Audit Run Metrics; LeaseFileAudit Runs library; Innovation Use Log",
            "notes": "Runs save_run and redirects to lease/property/portfolio view.",
        },
        {
            "action": "Bulk audit page",
            "routes": "GET /bulk-audit",
            "reads": "None (Entrata picklist + in-memory jobs)",
            "writes": "None",
            "notes": "No SharePoint persistence by this route itself.",
        },
        {
            "action": "Start bulk audit",
            "routes": "POST /api/bulk-audit",
            "reads": "None (Entrata picklist)",
            "writes": "None directly",
            "notes": "Creates in-memory job; worker does SharePoint writes.",
        },
        {
            "action": "Bulk audit worker (background)",
            "routes": "Triggered by /api/bulk-audit",
            "reads": "Entrata API",
            "writes": "AuditRuns2; RunDisplaySnapshots; Audit Run Metrics; LeaseFileAudit Runs library",
            "notes": "Per property save_run inside background thread.",
        },
        {
            "action": "Bulk audit status",
            "routes": "GET /api/bulk-audit/<job_id>",
            "reads": "None (in-memory jobs)",
            "writes": "None",
            "notes": "No SharePoint access.",
        },
        {
            "action": "Bulk audit cancel",
            "routes": "POST /api/bulk-audit/<job_id>/cancel",
            "reads": "None (in-memory jobs)",
            "writes": "None",
            "notes": "Signals cancellation only.",
        },
        {
            "action": "Portfolio view",
            "routes": "GET /portfolio and /portfolio/<run_id>",
            "reads": "RunDisplaySnapshots; ExceptionMonths; LeaseFileAudit Runs library",
            "writes": "None",
            "notes": "Aggregates property snapshots and exception summaries.",
        },
        {
            "action": "Property view",
            "routes": "GET /property/<property_id> and /property/<property_id>/<run_id>",
            "reads": "AuditRuns2; RunDisplaySnapshots; ExceptionMonths; LeaseFileAudit Runs library",
            "writes": "None",
            "notes": "Loads lease-level exception grouping for a property.",
        },
        {
            "action": "Lease view",
            "routes": "GET /lease/<property_id>/<lease_interval_id> and /lease/<run_id>/<property_id>/<lease_interval_id>",
            "reads": "AuditRuns2; RunDisplaySnapshots; ExceptionMonths; LeaseFileAudit Runs library",
            "writes": "None",
            "notes": "Detailed lease-level discrepancy display.",
        },
        {
            "action": "Bucket drilldown",
            "routes": "GET /bucket/<run_id>/<property_id>/<lease_interval_id>/<ar_code_id>/<audit_month>",
            "reads": "LeaseFileAudit Runs library (run files)",
            "writes": "None",
            "notes": "Reads run payload and renders bucket detail.",
        },
        {
            "action": "Get exception months",
            "routes": "GET /api/exception-months/<run_id>/<property_id>/<lease_interval_id>/<ar_code_id>",
            "reads": "ExceptionMonths",
            "writes": "None",
            "notes": "Returns month-level status rows for one AR code.",
        },
        {
            "action": "Upsert exception month",
            "routes": "POST /api/exception-months",
            "reads": "AuditRuns2 (for scoped status calc); ExceptionMonths",
            "writes": "ExceptionMonths",
            "notes": "Writes/updates month state and recalculates AR code status.",
        },
        {
            "action": "AR status lookup",
            "routes": "GET /api/exception-months/ar-status/<run_id>/<property_id>/<lease_interval_id>/<ar_code_id>",
            "reads": "AuditRuns2; ExceptionMonths",
            "writes": "None",
            "notes": "Computes scoped Open/Resolved status.",
        },
        {
            "action": "Lease terms API",
            "routes": "GET /api/lease-terms/<run_id>/<property_id>/<lease_interval_id>",
            "reads": "LeaseFileAudit Runs library; LeaseTerms (fallback read)",
            "writes": "LeaseTermSet; LeaseTerms; LeaseTermEvidence; LeaseFileAudit Runs library (lease PDFs), when refresh runs",
            "notes": "Calls refresh_lease_terms_for_lease_interval and returns expectation overlay.",
        },
    ]

    doc = Document()
    doc.add_heading("LeaseFileAudit SharePoint Action Matrix", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        "This matrix maps each app action/route to the SharePoint lists or libraries it reads from and writes to."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Action"
    header[1].text = "Route / Trigger"
    header[2].text = "Reads From SharePoint"
    header[3].text = "Writes To SharePoint"
    header[4].text = "Notes"

    for row in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], row["action"])
        _set_cell_text(cells[1], row["routes"])
        _set_cell_text(cells[2], row["reads"])
        _set_cell_text(cells[3], row["writes"])
        _set_cell_text(cells[4], row["notes"])

    doc.add_paragraph("")
    doc.add_paragraph("Current detailed results target is hard-locked to AuditRuns2 in storage/service.py.")

    doc.save(out_path)
    print(f"Created {out_path.resolve()}")


if __name__ == "__main__":
    main()
